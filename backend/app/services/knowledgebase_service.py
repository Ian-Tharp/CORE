from __future__ import annotations

from typing import Dict, List, Optional, Tuple
import asyncio
import os
import io
import logging
import math
import mimetypes
import shutil
import time

logger = logging.getLogger(__name__)

from app.repository import knowledgebase_repository as repo
from app.services.ollama_embeddings import embed_texts_via_ollama, embed_texts_batch


def _split_text(
    text: str, *, chunk_size: int = 1200, chunk_overlap: int = 200
) -> List[str]:
    if chunk_size <= 0:
        return [text]
    chunks: List[str] = []
    start = 0
    length = len(text)
    while start < length:
        end = min(start + chunk_size, length)
        chunks.append(text[start:end])
        if end == length:
            break
        start = end - chunk_overlap
        if start < 0:
            start = 0
    return chunks


def _cosine_similarity(vec_a: List[float], vec_b: List[float]) -> float:
    if not vec_a or not vec_b:
        return 0.0
    if len(vec_a) != len(vec_b):
        n = min(len(vec_a), len(vec_b))
        vec_a = vec_a[:n]
        vec_b = vec_b[:n]
    dot = sum(a * b for a, b in zip(vec_a, vec_b))
    norm_a = math.sqrt(sum(a * a for a in vec_a))
    norm_b = math.sqrt(sum(b * b for b in vec_b))
    if norm_a == 0 or norm_b == 0:
        return 0.0
    return dot / (norm_a * norm_b)


async def _embed_texts_local(
    *, model: str, texts: List[str], use_batch: bool = True
) -> Tuple[List[List[float]], int]:
    """Embed texts locally via Ollama and return (vectors, original_dim).

    When use_batch is True, uses the batch /api/embed endpoint (with fallback).
    """
    if use_batch:
        return await embed_texts_batch(model=model, texts=texts)
    return await embed_texts_via_ollama(model=model, texts=texts)


async def process_uploaded_file(
    *,
    storage_path: str,
    original_name: str,
    mime_type: str,
    description: Optional[str],
    is_global: bool,
    file_hash: Optional[str] = None,
    local_model: Optional[str] = None,
    instance_name: Optional[str] = None,
    source_discussion: Optional[str] = None,
) -> str:
    start_time = time.time()
    title, text = await _extract_title_and_text(storage_path, mime_type)
    model = local_model or "nomic-embed-text"

    # Auto-description: simple heuristic summary (first 1-2 sentences)
    auto_description: Optional[str] = None
    if not description:
        snippet = (text or "").strip()[:800]
        if snippet:
            parts = [p.strip() for p in snippet.split(".") if p.strip()]
            auto_description = ". ".join(parts[:2])[:200]

    # Create document entry
    stat = os.stat(storage_path)
    filename = os.path.basename(storage_path)
    doc_id = await repo.create_document(
        filename=filename,
        original_name=original_name,
        size=stat.st_size,
        mime_type=mime_type,
        storage_path=storage_path,
        description=description or auto_description,
        is_global=is_global,
        title=title,
        file_hash=file_hash,
        instance_name=instance_name,
        source_discussion=source_discussion,
    )

    # Document-level embedding
    title_desc = f"{(title or original_name)}\n\n{(description or auto_description or '')}".strip()
    if title_desc:
        vecs, orig = await _embed_texts_local(model=model, texts=[title_desc])
        if vecs:
            await repo.update_document_embedding(
                document_id=doc_id,
                embedding=vecs[0],
                model=model,
                dimensions=orig,
            )

    # Chunk text and embed
    chunks = _split_text(text)
    vecs_all: List[List[float]] = []
    dims = 0
    batch_size = 64
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start : start + batch_size]
        v, od = await _embed_texts_local(model=model, texts=batch)
        if not v:
            continue
        vecs_all.extend(v)
        if od and not dims:
            dims = od
    updates: List[Tuple[int, List[float]]] = []
    for idx, _ in enumerate(chunks):
        if idx < len(vecs_all):
            updates.append((idx, vecs_all[idx]))
    if updates:
        # Convert updates to inserts with text for new chunks
        chunk_inserts = []
        for idx, vec in updates:
            chunk_inserts.append((idx, chunks[idx], vec))
        await repo.insert_chunk_embeddings(
            document_id=doc_id,
            items=chunk_inserts,
            model=model,
            dimensions=dims or 0,
            instance_name=instance_name,
            source_discussion=source_discussion,
        )

    # Record document processing performance
    elapsed_ms = (time.time() - start_time) * 1000
    total_chars = len(text or "")
    try:
        from app.services.metrics_service import record_embedding_performance

        record_embedding_performance(
            "document", model, len(chunks), elapsed_ms, total_chars
        )
    except Exception:
        pass  # Don't fail on metrics errors

    return doc_id


async def ingest_world_wiki(
    world_id: str, *, local_model: Optional[str] = None
) -> Dict[str, int]:
    """Ingest a world's wiki pages into the knowledgebase as world-scoped docs.

    Each wiki page becomes a kb_documents row (source='wiki', world_id set) whose
    content is chunked and embedded for RAG. Pages already ingested (deduped by a
    ``wiki:<page_id>`` file_hash) are skipped, so this is safe to re-run. Embedding
    is best-effort: a document is still created/listed even if embedding fails.
    """
    from app.repository import creative_repository

    pages = await creative_repository.list_wiki_pages(world_id)
    model = local_model or os.getenv("EMBEDDING_MODEL") or "nomic-embed-text"
    ingested = 0
    skipped = 0

    for page in pages:
        page_id = page.get("id")
        file_hash = f"wiki:{page_id}"
        if await repo.get_document_by_hash(file_hash):
            skipped += 1
            continue

        title = page.get("title") or "Wiki page"
        content = page.get("content") or ""
        doc_id = await repo.create_document(
            filename=f"wiki-{page_id}.md",
            original_name=title,
            size=len(content.encode("utf-8")),
            mime_type="text/markdown",
            storage_path=f"wiki://{page_id}",
            description=(content.strip()[:200] or None),
            is_global=False,
            source="wiki",
            title=title,
            file_hash=file_hash,
            world_id=world_id,
        )

        # Best-effort chunk + embed — never let an embedding failure drop the doc.
        try:
            chunks = _split_text(content)
            vecs_all: List[List[float]] = []
            dims = 0
            for start in range(0, len(chunks), 64):
                v, od = await _embed_texts_local(
                    model=model, texts=chunks[start : start + 64]
                )
                if v:
                    vecs_all.extend(v)
                    if od and not dims:
                        dims = od
            items = [
                (i, chunks[i], vecs_all[i])
                for i in range(min(len(chunks), len(vecs_all)))
            ]
            if items:
                await repo.insert_chunk_embeddings(
                    document_id=doc_id, items=items, model=model, dimensions=dims or 0
                )
        except Exception as exc:  # noqa: BLE001
            logger.warning("Wiki embed failed for page %s: %s", page_id, exc)

        ingested += 1

    return {"ingested": ingested, "skipped": skipped, "total": len(pages)}


async def reprocess_document(
    *,
    document_id: str,
    storage_path: str,
    original_name: str,
    mime_type: str,
    description: Optional[str] = None,
    local_model: Optional[str] = None,
) -> None:
    """Re-extract text and re-embed an existing document (delete old chunks first)."""
    logger.info("Reprocessing document %s (%s)", document_id, original_name)

    # 1. Extract text with OCR fallback
    title, text = await _extract_title_and_text(storage_path, mime_type)
    logger.info(
        "Extracted %d chars, title=%r for %s", len(text or ""), title, document_id
    )

    # 2. Delete old chunks
    from app.dependencies import get_db_pool

    pool = await get_db_pool()
    async with pool.acquire() as conn:
        deleted = await conn.execute(
            "DELETE FROM kb_chunks WHERE document_id = $1", document_id
        )
        logger.info("Deleted old chunks for %s: %s", document_id, deleted)

    # 3. Update document title if extracted
    if title:
        async with pool.acquire() as conn:
            await conn.execute(
                "UPDATE kb_documents SET title = $1 WHERE id = $2", title, document_id
            )

    model = local_model or "nomic-embed-text"

    # 4. Document-level embedding
    title_desc = f"{(title or original_name)}\n\n{(description or '')}".strip()
    if title_desc:
        vecs, orig = await _embed_texts_local(model=model, texts=[title_desc])
        if vecs:
            await repo.update_document_embedding(
                document_id=document_id,
                embedding=vecs[0],
                model=model,
                dimensions=orig,
            )

    # 5. Chunk text and embed
    chunks = _split_text(text)
    logger.info("Split into %d chunks for %s", len(chunks), document_id)

    if not chunks:
        logger.warning(
            "No text chunks for %s — document may be empty or image-only", document_id
        )
        return

    vecs_all: List[List[float]] = []
    dims = 0
    batch_size = 64
    for start in range(0, len(chunks), batch_size):
        batch = chunks[start : start + batch_size]
        v, od = await _embed_texts_local(model=model, texts=batch)
        if v:
            vecs_all.extend(v)
            if od and not dims:
                dims = od
    # Insert new chunks with embeddings
    chunk_payload: List[Tuple[int, str, List[float]]] = []
    for idx, chunk_text in enumerate(chunks):
        if idx < len(vecs_all):
            chunk_payload.append((idx, chunk_text, vecs_all[idx]))
    if chunk_payload:
        await repo.insert_chunk_embeddings(
            document_id=document_id,
            items=chunk_payload,
            model=model,
            dimensions=dims or 0,
        )
    logger.info(
        "Reprocessing complete for %s: %d chunks embedded", document_id, len(chunks)
    )


async def embed_document_locally(*, document_id: str, model: str) -> None:
    """Generate and persist local embeddings for a document and its chunks via Ollama."""
    doc = await repo.get_document(document_id)
    if not doc:
        return
    chunks = await repo.list_chunks_for_document(document_id)

    # Document-level embedding using title/description fallback to first chunk
    title = doc.get("title") or doc.get("original_name") or doc.get("filename") or ""
    desc = doc.get("description") or ""
    title_desc = (f"{title}\n\n{desc}").strip()
    doc_texts: List[str] = [title_desc] if title_desc else []
    if not doc_texts and chunks:
        doc_texts = [chunks[0].get("text") or ""]
    if doc_texts:
        doc_vecs, original_dim = await _embed_texts_local(model=model, texts=doc_texts)
        if doc_vecs:
            await repo.update_document_embedding(
                document_id=document_id,
                embedding=doc_vecs[0],
                model=model,
                dimensions=original_dim,
            )

    # Chunk-level embeddings
    chunk_texts = [c.get("text") or "" for c in chunks]
    if chunk_texts:
        batch_size = 64
        items: List[Tuple[int, str, List[float]]] = []
        original_dim_total = 0
        for start in range(0, len(chunk_texts), batch_size):
            batch = chunk_texts[start : start + batch_size]
            vecs, original_dim = await _embed_texts_local(model=model, texts=batch)
            original_dim_total = original_dim or original_dim_total
            for i, v in enumerate(vecs):
                items.append((start + i, batch[i], v))
        if items:
            update_tuples = [(idx, vec) for idx, _text, vec in items]
            await repo.update_chunk_embeddings(
                document_id=document_id,
                chunks=update_tuples,
                model=model,
                dimensions=original_dim_total or 0,
            )


async def _extract_text(path: str, mime_type: str) -> str:
    # Handle simple text types
    if mime_type.startswith("text/") or mime_type in ("application/json",):
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()

    # Basic PDF support via pypdf if installed
    if mime_type == "application/pdf":
        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(path)
            text = []
            for page in reader.pages:
                try:
                    text.append(page.extract_text() or "")
                except Exception:
                    continue
            return "\n".join(text)
        except Exception:
            return ""

    # Basic DOCX support via python-docx if installed
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            import docx  # type: ignore

            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return ""

    # Fallback: try to guess and read as text
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


_OCR_MAX_PAGES = 500
_TESSERACT_AVAILABLE = shutil.which("tesseract") is not None
# Render scale for OCR rasterization (~144 DPI at the 72pt PDF base).
_OCR_RENDER_SCALE = 2.0


def _pdfium_page_text(page) -> str:
    """Extract the embedded text layer from a pypdfium2 page."""
    try:
        textpage = page.get_textpage()
    except Exception:
        return ""
    try:
        return textpage.get_text_range() or ""
    except Exception:
        return ""
    finally:
        try:
            textpage.close()
        except Exception:
            pass


def _ocr_pdfium_page(page) -> str:
    """Render a pypdfium2 page to an image and OCR it with pytesseract.

    Both pypdfium2 (Apache-2.0/BSD-3) and pytesseract (Apache-2.0) are
    permissively licensed, replacing the former AGPL pymupdf OCR path.
    """
    try:
        import pytesseract  # type: ignore

        bitmap = page.render(scale=_OCR_RENDER_SCALE)
        image = bitmap.to_pil()
        try:
            return pytesseract.image_to_string(image) or ""
        finally:
            try:
                image.close()
            except Exception:
                pass
    except Exception:
        logger.debug("pytesseract OCR failed", exc_info=True)
        return ""


def _extract_pdf_with_ocr(
    path: str, existing_title: Optional[str]
) -> Tuple[str, Optional[str]]:
    """Extract text from a PDF using pypdfium2 (text layer) with a pytesseract
    OCR fallback per page.

    Permissive-licensed replacement for the former pymupdf/fitz path (PyMuPDF
    is AGPL-3.0). pypdfium2 is fast and does not hang on malformed font maps,
    preserving the reason the pymupdf-first path originally existed.
    """
    try:
        import pypdfium2 as pdfium  # type: ignore
    except Exception:
        logger.warning(
            "pypdfium2 unavailable; cannot extract %s (install pypdfium2)",
            path,
            exc_info=True,
        )
        return "", existing_title

    title = existing_title
    pages_text: List[str] = []

    try:
        doc = pdfium.PdfDocument(path)
    except Exception:
        logger.error("pypdfium2 could not open %s", path, exc_info=True)
        return "", title

    try:
        total = len(doc)
        num_pages = min(total, _OCR_MAX_PAGES)
        if total > _OCR_MAX_PAGES:
            logger.warning(
                "PDF has %d pages; limiting OCR to first %d pages: %s",
                total,
                _OCR_MAX_PAGES,
                path,
            )

        for idx in range(num_pages):
            text = ""
            try:
                page = doc[idx]
                text = _pdfium_page_text(page) or ""
                if len(text.strip()) < 20 and _TESSERACT_AVAILABLE:
                    ocr_text = _ocr_pdfium_page(page)
                    if len(ocr_text.strip()) > len(text.strip()):
                        text = ocr_text
                        logger.debug("OCR used for page %d of %s", idx, path)
            except Exception:
                logger.debug(
                    "pypdfium2 failed on page %d of %s", idx, path, exc_info=True
                )
            pages_text.append(text)
    finally:
        try:
            doc.close()
        except Exception:
            pass

    full_text = "\n".join(pages_text)

    if not title and pages_text:
        for line in pages_text[0].splitlines():
            candidate = (line or "").strip()
            if 3 <= len(candidate) <= 140:
                title = candidate
                break

    logger.info(
        "pypdfium2 extracted %d chars from %d pages: %s",
        len(full_text),
        num_pages,
        path,
    )
    return full_text, title


async def _extract_title_and_text(
    path: str, mime_type: str, *, prefer_ocr: bool = False
) -> Tuple[Optional[str], str]:
    """Best-effort extraction of a human-friendly title and full text.

    If prefer_ocr=True, use the pypdfium2 path first (faster, no hanging on
    malformed fonts). Falls back to pypdf only if pypdfium2 fails.
    """
    title: Optional[str] = None

    # PDF path
    if mime_type == "application/pdf":
        if prefer_ocr:
            # pypdfium2-first path: faster and doesn't hang on font map issues
            try:
                text, fallback_title = _extract_pdf_with_ocr(path, None)
                if text and len(text.strip()) >= 100:
                    return fallback_title, text
                logger.info(
                    "pypdfium2 extracted < 100 chars for %s, trying pypdf", path
                )
            except Exception:
                logger.warning(
                    "pypdfium2 failed for %s, trying pypdf", path, exc_info=True
                )

        try:
            from pypdf import PdfReader  # type: ignore

            reader = PdfReader(path)
            pages_text: List[str] = []
            first_page_text: str = ""
            for idx, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text() or ""
                    pages_text.append(page_text)
                    if idx == 0:
                        first_page_text = page_text
                except Exception:
                    continue

            try:
                meta = getattr(reader, "metadata", None)
                meta_title = None
                if meta is not None:
                    meta_title = getattr(meta, "title", None)
                if isinstance(meta_title, str):
                    cleaned = meta_title.strip()
                    title = cleaned if cleaned else None
            except Exception:
                title = None

            if not title and first_page_text:
                for line in first_page_text.splitlines() or []:
                    candidate = (line or "").strip()
                    if 3 <= len(candidate) <= 140:
                        title = candidate
                        break

            full_text = "\n".join(pages_text)

            num_pages = len(reader.pages)
            if len(full_text.strip()) < 100 and num_pages > 0:
                logger.info(
                    "pypdf extracted < 100 chars from %d-page PDF, "
                    "falling back to pypdfium2 OCR: %s",
                    num_pages,
                    path,
                )
                full_text, title = _extract_pdf_with_ocr(path, title)
            else:
                logger.debug("pypdf text extraction succeeded for %s", path)

            return title, full_text
        except Exception:
            logger.warning("pypdf failed for %s, trying pypdfium2", path, exc_info=True)
            try:
                text, fallback_title = _extract_pdf_with_ocr(path, None)
                return fallback_title, text
            except Exception:
                logger.warning("pypdfium2 also failed for %s", path, exc_info=True)
            pass

    # DOCX path
    if mime_type in (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ):
        try:
            import docx  # type: ignore

            doc = docx.Document(path)
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
            for p in paragraphs:
                candidate = (p or "").strip()
                if candidate:
                    title = candidate[:140]
                    break
            return title, text
        except Exception:
            pass

    # Plain text and everything else
    try:
        with open(path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()
        if content:
            for line in content.splitlines():
                candidate = (line or "").strip()
                if candidate:
                    title = candidate[:140]
                    break
        return title, content
    except Exception:
        return None, ""


async def reextract_title_for_document(
    *, storage_path: str, mime_type: str
) -> Optional[str]:
    """Re-extract a best-effort title from the stored document without re-embedding."""
    title, _ = await _extract_title_and_text(storage_path, mime_type)
    return title


async def retrieve_context(
    *,
    query: str,
    mode: str,
    file_id: Optional[str] = None,
    max_docs: int = 5,
    max_chunks: int = 8,
    local_model: Optional[str] = None,
) -> Dict[str, any]:
    model = (local_model or "nomic-embed-text").strip()
    qvecs, _orig = await _embed_texts_local(model=model, texts=[query])
    if not qvecs:
        return {"chunks": [], "doc_ids": []}
    query_vec = qvecs[0]

    # Determine candidate documents
    candidate_doc_ids: List[str] = []
    if mode == "file" and file_id:
        candidate_doc_ids = [file_id]

    # SQL vector search across local vectors
    doc_filter = candidate_doc_ids if mode == "file" else None
    rows = await repo.search_chunks_by_vector(
        query_vec=query_vec,
        limit=max_chunks,
        document_filter=doc_filter,
        model=local_model,
    )
    return {"chunks": rows, "doc_ids": list({r.get("document_id") for r in rows})}


async def retrieve_context_by_instance(
    *,
    query: str,
    instance_name: str,
    max_chunks: int = 8,
    local_model: Optional[str] = None,
    document_filter: Optional[List[str]] = None,
) -> Dict[str, any]:
    """Retrieve context filtered by instance perspective for 'what does [instance] think about X?' queries."""
    model = (local_model or "nomic-embed-text").strip()
    qvecs, _orig = await _embed_texts_local(model=model, texts=[query])
    if not qvecs:
        return {"chunks": [], "doc_ids": [], "instance_name": instance_name}
    query_vec = qvecs[0]

    rows = await repo.search_chunks_by_instance(
        query_vec=query_vec,
        instance_name=instance_name,
        limit=max_chunks,
        document_filter=document_filter,
        model=local_model,
    )
    return {
        "chunks": rows,
        "doc_ids": list({r.get("document_id") for r in rows}),
        "instance_name": instance_name,
    }


async def retrieve_context_by_world(
    *,
    query: str,
    world_id: str,
    max_chunks: int = 8,
    local_model: Optional[str] = None,
) -> Dict[str, any]:
    """Retrieve RAG context limited to a single world's ingested knowledge.

    Scopes the vector search to the world's documents so an agent reasoning about
    a world only pulls that world's lore.
    """
    docs = await repo.list_documents_by_world(world_id)
    doc_ids = [d["id"] for d in docs]
    if not doc_ids:
        return {"results": [], "doc_ids": [], "world_id": world_id}

    model = (local_model or os.getenv("EMBEDDING_MODEL") or "nomic-embed-text").strip()
    qvecs, _orig = await _embed_texts_local(model=model, texts=[query])
    if not qvecs:
        return {"results": [], "doc_ids": [], "world_id": world_id}

    rows = await repo.search_chunks_by_vector(
        query_vec=qvecs[0],
        limit=max_chunks,
        document_filter=doc_ids,
        model=local_model,
    )
    return {
        "results": rows,
        "doc_ids": list({r.get("document_id") for r in rows}),
        "world_id": world_id,
    }


async def batch_upload_and_process(
    *,
    file_paths: List[str],
    is_global: bool = False,
    local_model: Optional[str] = None,
    description: Optional[str] = None,
) -> Dict:
    """Process multiple files: extract text, chunk, batch-embed, store.

    Returns dict with doc_ids, stats per doc, and totals.
    """
    import time
    import hashlib

    model = local_model or "nomic-embed-text"
    start_time = time.monotonic()

    # Phase 1: Extract and chunk all files (with per-file timeout)
    file_infos: List[Dict] = []  # {path, title, text, chunks, error}
    extraction_timeout = 120  # seconds per file
    for i, fpath in enumerate(file_paths):
        logger.info("Processing file %d of %d: %s", i + 1, len(file_paths), fpath)
        try:
            mime = mimetypes.guess_type(fpath)[0] or "application/octet-stream"
            # Run extraction with timeout — pypdf can hang on malformed PDFs
            title, text = await asyncio.wait_for(
                _extract_title_and_text(fpath, mime, prefer_ocr=True),
                timeout=extraction_timeout,
            )
            # Strip null bytes — some PDFs contain \x00 that PostgreSQL rejects
            if text:
                text = text.replace("\x00", "")
            if title:
                title = title.replace("\x00", "")
            chunks = _split_text(text) if text else []
            file_infos.append(
                {
                    "path": fpath,
                    "title": title,
                    "text": text,
                    "chunks": chunks,
                    "mime": mime,
                    "error": None,
                }
            )
        except asyncio.TimeoutError:
            logger.error(
                "Extraction timed out after %ds: %s", extraction_timeout, fpath
            )
            file_infos.append(
                {
                    "path": fpath,
                    "chunks": [],
                    "error": f"extraction timed out after {extraction_timeout}s",
                }
            )
        except Exception as exc:
            logger.error("Failed to extract %s: %s", fpath, exc)
            file_infos.append({"path": fpath, "chunks": [], "error": str(exc)})

    # Phase 2: Collect all chunks and batch embed
    all_chunks: List[str] = []
    chunk_map: List[Tuple[int, int]] = []  # (file_index, chunk_index_within_file)
    for fi, info in enumerate(file_infos):
        if info.get("error"):
            continue
        for ci, chunk in enumerate(info["chunks"]):
            all_chunks.append(chunk)
            chunk_map.append((fi, ci))

    all_vecs: List[List[float]] = []
    dims = 0
    batch_size = 64
    total_batches = (
        (len(all_chunks) + batch_size - 1) // batch_size if all_chunks else 0
    )
    for batch_idx, start in enumerate(range(0, len(all_chunks), batch_size)):
        logger.info("Embedding batch %d of %d", batch_idx + 1, total_batches)
        batch = all_chunks[start : start + batch_size]
        v, od = await embed_texts_batch(model=model, texts=batch)
        if v:
            all_vecs.extend(v)
            if od and not dims:
                dims = od

    # Phase 3: Store documents and chunks
    doc_ids: List[str] = []
    stats: List[Dict] = []
    vec_offset = 0

    for fi, info in enumerate(file_infos):
        if info.get("error"):
            stats.append({"file": info["path"], "error": info["error"], "chunks": 0})
            continue

        fpath = info["path"]
        try:
            stat = os.stat(fpath)
            filename = os.path.basename(fpath)
            with open(fpath, "rb") as f:
                file_hash = hashlib.sha256(f.read()).hexdigest()

            auto_desc = description
            if not auto_desc:
                snippet = (info.get("text") or "").strip()[:800]
                if snippet:
                    parts = [p.strip() for p in snippet.split(".") if p.strip()]
                    auto_desc = ". ".join(parts[:2])[:200]

            doc_id = await repo.create_document(
                filename=filename,
                original_name=filename,
                size=stat.st_size,
                mime_type=info["mime"],
                storage_path=fpath,
                description=auto_desc,
                is_global=is_global,
                title=info.get("title"),
                file_hash=file_hash,
            )

            # Document-level embedding
            title_desc = (
                f"{(info.get('title') or filename)}\n\n{(auto_desc or '')}".strip()
            )
            if title_desc:
                dvecs, dorig = await embed_texts_batch(model=model, texts=[title_desc])
                if dvecs:
                    await repo.update_document_embedding(
                        document_id=doc_id,
                        embedding=dvecs[0],
                        model=model,
                        dimensions=dorig,
                    )

            # Chunk embeddings — INSERT (not update, chunks don't exist yet)
            num_chunks = len(info["chunks"])
            inserts: List[Tuple[int, str, List[float]]] = []
            for ci in range(num_chunks):
                if vec_offset + ci < len(all_vecs):
                    inserts.append((ci, info["chunks"][ci], all_vecs[vec_offset + ci]))
            if inserts:
                await repo.insert_chunk_embeddings(
                    document_id=doc_id,
                    items=inserts,
                    model=model,
                    dimensions=dims or 0,
                )
            vec_offset += num_chunks

            doc_ids.append(doc_id)
            stats.append({"file": fpath, "doc_id": doc_id, "chunks": num_chunks})
        except Exception as exc:
            logger.error("Failed to store %s: %s", fpath, exc)
            stats.append({"file": fpath, "error": str(exc), "chunks": 0})
            vec_offset += len(info["chunks"])

    elapsed = time.monotonic() - start_time
    return {
        "doc_ids": doc_ids,
        "stats": stats,
        "total_chunks": len(all_chunks),
        "total_time": round(elapsed, 2),
    }


def build_rag_messages(
    original_messages: List[Dict[str, str]], *, context_chunks: List[Dict[str, any]]
) -> List[Dict[str, str]]:
    if not context_chunks:
        return original_messages
    ctx_lines = []
    for ch in context_chunks:
        ref = f"doc:{ch['document_id']}#chunk:{ch['chunk_index']}"
        snippet = ch["text"]
        ctx_lines.append(f"[{ref}]\n{snippet}")
    context_block = "\n\n".join(ctx_lines)
    system_msg = {
        "role": "system",
        "content": (
            "You are a precise assistant. Use the provided knowledgebase context when relevant. "
            "If the context is insufficient, say so and answer from your general knowledge.\n\n"
            f"Context:\n{context_block}"
        ),
    }
    return [system_msg] + original_messages


# =============================================================================
# KNOWLEDGE ATTRIBUTION VERIFICATION
# =============================================================================


async def verify_knowledge_attribution(*, chunk_ids: List[str]) -> Dict[str, any]:
    """
    Verify knowledge attribution for given chunk IDs.

    Checks if attribution metadata exists and source instance is valid.
    Returns verification report with warnings for orphaned attributions.

    Args:
        chunk_ids: List of chunk IDs to verify attribution for

    Returns:
        Dict with verification results and any warnings
    """
    if not chunk_ids:
        return {"verified": [], "warnings": [], "total_chunks": 0}

    logger.info("Verifying attribution for %d chunks", len(chunk_ids))

    # Get chunk attribution data
    chunk_attributions = await repo.get_chunk_attributions(chunk_ids)

    verified_chunks = []
    warnings = []

    for chunk_id in chunk_ids:
        chunk_attribution = chunk_attributions.get(chunk_id)

        if not chunk_attribution:
            warnings.append(
                {
                    "type": "missing_attribution",
                    "chunk_id": chunk_id,
                    "message": f"Chunk {chunk_id} has no attribution metadata",
                }
            )
            continue

        # Check if source instance exists and is valid
        source_instance_id = chunk_attribution.get("source_instance_id")
        if source_instance_id:
            from app.repository import instance_repository

            instance = await instance_repository.get_instance(source_instance_id)

            if not instance:
                warnings.append(
                    {
                        "type": "orphaned_attribution",
                        "chunk_id": chunk_id,
                        "source_instance_id": source_instance_id,
                        "message": f"Chunk {chunk_id} attributes to non-existent instance {source_instance_id}",
                    }
                )
            else:
                verified_chunks.append(
                    {
                        "chunk_id": chunk_id,
                        "attribution": chunk_attribution,
                        "source_instance": {
                            "id": str(instance.id),
                            "agent_id": instance.agent_id,
                            "agent_role": instance.agent_role,
                            "created_at": instance.created_at.isoformat(),
                        },
                    }
                )
        else:
            warnings.append(
                {
                    "type": "missing_source_instance",
                    "chunk_id": chunk_id,
                    "message": f"Chunk {chunk_id} has attribution but no source_instance_id",
                }
            )

    return {
        "verified": verified_chunks,
        "warnings": warnings,
        "total_chunks": len(chunk_ids),
        "verified_count": len(verified_chunks),
        "warning_count": len(warnings),
    }


async def verify_document_knowledge_integrity(*, document_id: str) -> Dict[str, any]:
    """
    Verify knowledge integrity for all chunks in a document.

    Comprehensive verification of attribution metadata, instance validity,
    and knowledge graph consistency for a single document.

    Args:
        document_id: Document ID to verify

    Returns:
        Dict with detailed integrity report
    """
    logger.info("Verifying knowledge integrity for document %s", document_id)

    # Get document info
    document = await repo.get_document(document_id)
    if not document:
        return {"error": f"Document {document_id} not found", "verified": False}

    # Get all chunks for the document
    chunks = await repo.list_chunks_for_document(document_id)
    chunk_ids = [chunk["id"] for chunk in chunks]

    if not chunk_ids:
        return {
            "document_id": document_id,
            "verified": True,
            "chunks_count": 0,
            "message": "Document has no chunks to verify",
        }

    # Verify attribution for all chunks
    attribution_report = await verify_knowledge_attribution(chunk_ids=chunk_ids)

    # Additional document-level checks
    doc_warnings = []

    # Check document-level attribution if it exists
    doc_attribution = await repo.get_document_attribution(document_id)
    if doc_attribution:
        source_instance_id = doc_attribution.get("source_instance_id")
        if source_instance_id:
            from app.repository import instance_repository

            instance = await instance_repository.get_instance(source_instance_id)
            if not instance:
                doc_warnings.append(
                    {
                        "type": "orphaned_document_attribution",
                        "document_id": document_id,
                        "source_instance_id": source_instance_id,
                        "message": f"Document {document_id} attributes to non-existent instance {source_instance_id}",
                    }
                )

    return {
        "document_id": document_id,
        "document_title": document.get("title") or document.get("filename", "Unknown"),
        "verified": len(attribution_report["warnings"]) == 0 and len(doc_warnings) == 0,
        "chunks_count": len(chunks),
        "chunk_attribution_report": attribution_report,
        "document_warnings": doc_warnings,
        "summary": {
            "total_issues": len(attribution_report["warnings"]) + len(doc_warnings),
            "chunk_issues": len(attribution_report["warnings"]),
            "document_issues": len(doc_warnings),
        },
    }


async def batch_verify_knowledge_attribution(
    *, limit: Optional[int] = None
) -> Dict[str, any]:
    """
    Batch verification of knowledge attribution across the entire knowledge base.

    Scans all documents and chunks to identify attribution issues.
    Useful for maintenance and integrity monitoring.

    Args:
        limit: Optional limit on number of documents to process

    Returns:
        Dict with comprehensive attribution health report
    """
    logger.info("Starting batch knowledge attribution verification (limit: %s)", limit)

    # Get all documents
    documents = await repo.list_documents()
    if limit:
        documents = documents[:limit]

    total_documents = len(documents)
    verified_documents = []
    failed_verifications = []
    total_warnings = 0

    for i, doc in enumerate(documents):
        try:
            logger.debug(
                "Verifying document %d/%d: %s", i + 1, total_documents, doc["id"]
            )

            verification_result = await verify_document_knowledge_integrity(
                document_id=doc["id"]
            )

            if verification_result.get("error"):
                failed_verifications.append(
                    {"document_id": doc["id"], "error": verification_result["error"]}
                )
            else:
                verified_documents.append(verification_result)
                total_warnings += verification_result["summary"]["total_issues"]

        except Exception as exc:
            logger.error("Failed to verify document %s: %s", doc["id"], exc)
            failed_verifications.append({"document_id": doc["id"], "error": str(exc)})

    # Generate summary statistics
    healthy_documents = sum(1 for doc in verified_documents if doc["verified"])

    return {
        "batch_verification_complete": True,
        "processed_documents": total_documents,
        "verified_documents": len(verified_documents),
        "failed_verifications": len(failed_verifications),
        "healthy_documents": healthy_documents,
        "documents_with_issues": len(verified_documents) - healthy_documents,
        "total_attribution_warnings": total_warnings,
        "health_score": (
            healthy_documents / total_documents if total_documents > 0 else 1.0
        ),
        "detailed_results": verified_documents,
        "failures": failed_verifications,
    }
